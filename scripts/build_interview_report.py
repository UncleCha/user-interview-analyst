"""
用户访谈分析报告 Word 生成脚本
用法:
python scripts/build_interview_report.py --content-file report.json --output report.docx
"""

import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("需要安装 python-docx: pip install python-docx")
    sys.exit(1)


FONT_ZH = "微软雅黑"
FONT_EN = "Microsoft YaHei"


def set_run_font(run, size=11, bold=False, color_hex=None):
    """设置 run 字体，确保中英文一致。"""
    run.font.name = FONT_ZH
    run.font.size = Pt(size)
    run.font.bold = bold

    if color_hex:
        r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)

    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_EN)
    rFonts.set(qn("w:eastAsia"), FONT_ZH)
    rFonts.set(qn("w:hAnsi"), FONT_EN)
    rFonts.set(qn("w:cs"), FONT_ZH)

    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def set_cell_background(cell, color_hex):
    """设置表格单元格背景色。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def write_inline_bold(paragraph, text, font_size=11):
    """支持 **加粗** 的行内文本写入。"""
    parts = text.split("**")
    for idx, part in enumerate(parts):
        run = paragraph.add_run(part)
        set_run_font(run, size=font_size, bold=(idx % 2 == 1))


def normalize_table_row(line):
    raw = line.strip()
    if not raw:
        return []
    if raw.startswith("|"):
        raw = raw[1:]
    if raw.endswith("|"):
        raw = raw[:-1]
    return [cell.strip() for cell in raw.split("|")]


def is_table_separator_row(cells):
    if not cells:
        return False
    return all(re.match(r"^:?-{3,}:?$", cell) for cell in cells)


def parse_table_block(lines, start_index):
    """从 start_index 开始解析 markdown 表格。"""
    block = []
    idx = start_index
    while idx < len(lines):
        line = lines[idx]
        if "|" not in line or not line.strip():
            break
        block.append(line)
        idx += 1

    if len(block) < 2:
        return None, start_index

    header = normalize_table_row(block[0])
    divider = normalize_table_row(block[1])
    if not header or not is_table_separator_row(divider):
        return None, start_index

    data_rows = []
    for row_line in block[2:]:
        row_cells = normalize_table_row(row_line)
        if row_cells:
            data_rows.append(row_cells)

    return {"headers": header, "rows": data_rows}, idx


def render_table(doc, headers, rows):
    if not headers:
        return

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        run = cell.paragraphs[0].runs[0]
        set_run_font(run, size=10, bold=True, color_hex="FFFFFF")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, "1F2937")

    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx in range(len(headers)):
            text = row_data[c_idx] if c_idx < len(row_data) else ""
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = text
            if cell.paragraphs and cell.paragraphs[0].runs:
                set_run_font(cell.paragraphs[0].runs[0], size=10)
            set_cell_background(cell, "F9FAFB" if r_idx % 2 == 0 else "FFFFFF")

    doc.add_paragraph()


def add_bullet_paragraph(doc, text):
    try:
        para = doc.add_paragraph(style="List Bullet")
        write_inline_bold(para, text, font_size=11)
    except Exception:
        para = doc.add_paragraph()
        run = para.add_run("• ")
        set_run_font(run, size=11)
        write_inline_bold(para, text, font_size=11)


def render_markdown_content(doc, content):
    lines = content.splitlines()
    idx = 0

    while idx < len(lines):
        raw_line = lines[idx]
        line = raw_line.strip()

        if not line:
            idx += 1
            continue

        # 表格
        if "|" in line:
            table_data, next_idx = parse_table_block(lines, idx)
            if table_data is not None:
                render_table(doc, table_data["headers"], table_data["rows"])
                idx = next_idx
                continue

        # 子标题
        if line.startswith("### "):
            heading = doc.add_heading(line[4:].strip(), level=2)
            if heading.runs:
                set_run_font(heading.runs[0], size=13, bold=True, color_hex="111827")
            idx += 1
            continue

        # 列表
        if line.startswith("- "):
            add_bullet_paragraph(doc, line[2:].strip())
            idx += 1
            continue

        # 普通段落
        para = doc.add_paragraph()
        write_inline_bold(para, line, font_size=11)
        idx += 1


def add_cover(doc, title, subtitle):
    doc.add_paragraph()
    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    set_run_font(title_run, size=26, bold=True, color_hex="111827")

    if subtitle:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run(subtitle)
        set_run_font(sub_run, size=12, color_hex="6B7280")

    doc.add_page_break()


def add_section(doc, heading, content):
    h = doc.add_heading(heading, level=1)
    if h.runs:
        set_run_font(h.runs[0], size=15, bold=True, color_hex="111827")

    if content:
        render_markdown_content(doc, content)


def create_document(content_data, output_path):
    doc = Document()

    # 页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = content_data.get("title", "用户访谈结构化分析报告")
    subtitle = content_data.get("subtitle", "")
    sections = content_data.get("sections", [])
    notes = content_data.get("notes", "")

    add_cover(doc, title, subtitle)

    for section_data in sections:
        heading = section_data.get("heading", "未命名章节")
        content = section_data.get("content", "")
        add_section(doc, heading, content)

    if notes:
        add_section(doc, "备注", notes)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"[OK] 已生成: {output_path}")


def main():
    parser = argparse.ArgumentParser(description="生成用户访谈分析 Word 报告")
    parser.add_argument("--content-file", required=True, help="内容 JSON 文件路径")
    parser.add_argument("--output", required=True, help="输出 .docx 文件路径")
    args = parser.parse_args()

    content_path = Path(args.content_file)
    if not content_path.exists():
        print(f"内容文件不存在: {content_path}")
        sys.exit(1)

    with open(content_path, "r", encoding="utf-8") as f:
        content_data = json.load(f)

    output_path = Path(args.output)
    create_document(content_data, output_path)


if __name__ == "__main__":
    main()
